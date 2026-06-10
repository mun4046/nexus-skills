# -*- coding: utf-8 -*-
"""
매수콜(매매전략) 보고서 생성기.
build(D) 에 종목별 dict 를 넘기면 골드 브랜딩 Word 보고서(.docx)를 생성한다.

경로 설정(환경변수, 선택):
  BUYCALL_LOGO : 로고 PNG 경로 (기본: ../assets/nexus_logo.png)
  BUYCALL_OUT  : 출력 폴더    (기본: ./reports)

CLI 사용:
  python gen_report.py stocks.json            # JSON 배열의 각 종목 dict 빌드
  python gen_report.py stocks.json D:/out     # 출력 폴더 지정
JSON 의 tpcalc 항목은 [["값","강조여부(true/false)"], ...] 형태(파이썬 튜플 대응).

Windows 한글 깨짐 방지: PYTHONUTF8=1 python -X utf8 gen_report.py ...
"""
import sys, io, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- 브랜드 컬러 (회사에 맞게 교체 가능) ----
GOLD='9D844E'; GOLDD='7A663C'; LGOLD='F1EBDD'; LIGHT='F5F2EA'
DGRAY='333333'; GRAY='808080'; LINE='C9BFA8'; KFONT='맑은 고딕'

HERE=os.path.dirname(os.path.abspath(__file__))
LOGO=os.environ.get('BUYCALL_LOGO', os.path.normpath(os.path.join(HERE,'..','assets','nexus_logo.png')))
OUTDIR=os.environ.get('BUYCALL_OUT', os.path.join(os.getcwd(),'reports'))

def setfont(run,size=None,bold=None,italic=None,color=None,name=KFONT):
    run.font.name=name; rPr=run._element.get_or_add_rPr()
    rF=rPr.find(qn('w:rFonts'))
    if rF is None: rF=OxmlElement('w:rFonts'); rPr.append(rF)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rF.set(qn(a),name)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.font.bold=bold
    if italic is not None: run.font.italic=italic
    if color is not None: run.font.color.rgb=RGBColor.from_string(color)
def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color); tcPr.append(shd)
def vcenter(cell): cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
def cell_text(cell,text,size=9,bold=False,color=None,align='left',italic=False,space=0):
    cell.text=''; p=cell.paragraphs[0]
    p.alignment={'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after=Pt(space); p.paragraph_format.space_before=Pt(space); p.paragraph_format.line_spacing=1.0
    for i,seg in enumerate(str(text).split('\n')):
        if i: p.add_run().add_break()
        r=p.add_run(seg); setfont(r,size=size,bold=bold,color=color,italic=italic)
    vcenter(cell); return cell
def borders(table,color=LINE,sz=4,val='single'):
    tblPr=table._tbl.tblPr; old=tblPr.find(qn('w:tblBorders'))
    if old is not None: tblPr.remove(old)
    b=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),val)
        if val!='nil':
            e.set(qn('w:sz'),str(sz)); e.set(qn('w:space'),'0'); e.set(qn('w:color'),color)
        b.append(e)
    tblPr.append(b)
def colw(table,widths):
    table.autofit=False
    for row in table.rows:
        for i,w in enumerate(widths): row.cells[i].width=Cm(w)
def row_height(row,cm,rule='atLeast'):
    trPr=row._tr.get_or_add_trPr(); h=OxmlElement('w:trHeight')
    h.set(qn('w:val'),str(int(cm*567))); h.set(qn('w:hRule'),rule); trPr.append(h)
def underline_bar(p,color=GOLD,sz=14):
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bo=OxmlElement('w:bottom')
    bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),str(sz)); bo.set(qn('w:space'),'3'); bo.set(qn('w:color'),color)
    pbdr.append(bo); pPr.append(pbdr)
def num_heading(doc,n,ko,en='',space_before=10,pb=False):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(space_before); p.paragraph_format.space_after=Pt(4)
    if pb: p.paragraph_format.page_break_before=True
    setfont(p.add_run(f'{n}   '),size=13,bold=True,color=GOLD)
    setfont(p.add_run(ko),size=12,bold=True,color=GOLDD)
    if en: setfont(p.add_run(f'   ({en})'),size=9,bold=False,color=GRAY)
    underline_bar(p); return p
def heading_bar(doc,text,space_before=10,size=12,pb=False):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(space_before); p.paragraph_format.space_after=Pt(4)
    if pb: p.paragraph_format.page_break_before=True
    setfont(p.add_run(text),size=size,bold=True,color=GOLDD); underline_bar(p); return p
def subhead(doc,text,space_before=6):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(space_before); p.paragraph_format.space_after=Pt(3)
    setfont(p.add_run(text),size=10,bold=True,color=GOLDD); return p
def body(doc,text,size=9.5,color=None,italic=False,bullet=False,space=3,ls=1.18,lead=None,lead_color=None):
    p=doc.add_paragraph(style='List Bullet' if bullet else None)
    p.paragraph_format.space_after=Pt(space); p.paragraph_format.line_spacing=ls
    if lead: setfont(p.add_run(lead),size=size,bold=True,color=lead_color or GOLDD)
    setfont(p.add_run(text),size=size,color=color,italic=italic); return p
def sb_box(cell,title,rows,widths=(2.55,2.55),val_size=8,first=False):
    if first:
        cell.paragraphs[0].paragraph_format.space_after=Pt(0); setfont(cell.paragraphs[0].add_run(''),size=3)
    t=cell.add_table(rows=len(rows)+1,cols=2); borders(t,LINE,3); colw(t,list(widths))
    t.rows[0].cells[0].merge(t.rows[0].cells[1]); shade(t.rows[0].cells[0],GOLDD)
    cell_text(t.rows[0].cells[0],title,size=8.5,bold=True,color='FFFFFF',align='center')
    for i,(l,v) in enumerate(rows,start=1):
        shade(t.rows[i].cells[0],LIGHT)
        cell_text(t.rows[i].cells[0],l,size=7.8,bold=True,color=DGRAY)
        cell_text(t.rows[i].cells[1],v,size=val_size,align='right')
    sp=cell.add_paragraph(); sp.paragraph_format.space_after=Pt(2); sp.paragraph_format.space_before=Pt(0)
    setfont(sp.add_run(''),size=3); return t

def build(D, outdir=None):
    outdir=outdir or OUTDIR
    os.makedirs(outdir, exist_ok=True)
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Cm(1.2); sec.bottom_margin=Cm(1.2); sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)
    nm=doc.styles['Normal']; nm.font.name=KFONT; nm.font.size=Pt(9.5)
    nm.element.rPr.rFonts.set(qn('w:eastAsia'),KFONT)
    nm.paragraph_format.space_after=Pt(0); nm.paragraph_format.space_before=Pt(0); nm.paragraph_format.line_spacing=1.0
    hd=doc.add_table(rows=1,cols=2); hd.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(hd,val='nil'); colw(hd,[9.0,9.0])
    lc=hd.rows[0].cells[0]; lc.text=''; lp=lc.paragraphs[0]; lp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_after=Pt(0)
    if os.path.exists(LOGO): lp.add_run().add_picture(LOGO,width=Cm(3.2))
    else: setfont(lp.add_run('NEXUS ASSET'),size=14,bold=True,color=GOLDD)
    vcenter(lc)
    cell_text(hd.rows[0].cells[1],f"매매전략 Note   |   {D['date']}",size=9,color=GRAY,align='right')
    rule=doc.add_paragraph(); rule.paragraph_format.space_before=Pt(0); rule.paragraph_format.space_after=Pt(1)
    underline_bar(rule,color=GOLDD,sz=18)
    outer=doc.add_table(rows=1,cols=2); outer.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(outer,val='nil'); colw(outer,[5.7,12.1])
    side=outer.rows[0].cells[0]; main=outer.rows[0].cells[1]
    side.vertical_alignment=WD_ALIGN_VERTICAL.TOP; main.vertical_alignment=WD_ALIGN_VERTICAL.TOP; side.text=''; main.text=''
    sb_box(side,'투자의견',[('투자의견',D['opinion']),('목표주가',D['tp']),('현재주가',D['cp']),('상승여력',D['upside'])],val_size=9,first=True)
    sb_box(side,'Stock Data',D['stock'])
    rt=side.add_table(rows=2,cols=5); borders(rt,LINE,3); colw(rt,[1.4,1.0,1.0,1.0,1.1])
    rt.rows[0].cells[0].merge(rt.rows[0].cells[4]); shade(rt.rows[0].cells[0],GOLDD)
    cell_text(rt.rows[0].cells[0],'상대수익률 (KOSPI 대비, %p)',size=8,bold=True,color='FFFFFF',align='center')
    for i,h in enumerate(['','1M','3M','6M','12M']):
        shade(rt.rows[1].cells[i],LIGHT if i==0 else 'FFFFFF'); cell_text(rt.rows[1].cells[i],h,size=7.5,bold=True,color=DGRAY,align='center')
    vrow=side.add_table(rows=1,cols=5); borders(vrow,LINE,3); colw(vrow,[1.4,1.0,1.0,1.0,1.1])
    for i,v in enumerate(['초과']+list(D['rel'])):
        cell_text(vrow.rows[0].cells[i],v,size=7.6,bold=(i==0),color=DGRAY if i==0 else '000000',align='center')
    sp=side.add_paragraph(); sp.paragraph_format.space_after=Pt(2); setfont(sp.add_run(''),size=3)
    ch=side.add_table(rows=2,cols=1); borders(ch,LINE,3); colw(ch,[5.1])
    shade(ch.rows[0].cells[0],GOLDD); cell_text(ch.rows[0].cells[0],'주가 추이 (12M)',size=8.5,bold=True,color='FFFFFF',align='center')
    cell_text(ch.rows[1].cells[0],'[ 주가/상대강도 차트 삽입 ]',size=8,color=GRAY,align='center'); row_height(ch.rows[1],1.2)
    sp=side.add_paragraph(); sp.paragraph_format.space_after=Pt(2); setfont(sp.add_run(''),size=3)
    rep=D.get('reporter',['Reported by','○○○  (PM / Buy-side)','○○○@nexusasset.co.kr'])
    for j,txt in enumerate(rep):
        ap=side.add_paragraph(); ap.paragraph_format.space_before=Pt(0); ap.paragraph_format.space_after=Pt(0)
        setfont(ap.add_run(txt),size=7.5 if j!=1 else 8.5,bold=(j==1),color=GOLDD if j==1 else GRAY)
    p=main.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    setfont(p.add_run(D['name']+' '),size=17,bold=True,color=GOLDD); setfont(p.add_run('('+D['code']+')'),size=12,bold=True,color=DGRAY)
    p2=main.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(4)
    setfont(p2.add_run(D['sector']),size=8.5,color=GRAY)
    ph=main.add_paragraph(); ph.paragraph_format.space_after=Pt(4)
    setfont(ph.add_run(D['headline']),size=13,bold=True,color='000000')
    tpt=main.add_paragraph(); tpt.paragraph_format.space_after=Pt(2)
    setfont(tpt.add_run('▌ 매매전략'),size=10,bold=True,color=GOLDD)
    tp=main.add_table(rows=3,cols=4); borders(tp,LINE,4); colw(tp,[2.7,3.35,2.7,3.35])
    plan=[('전략 유형',D['stype'],'투자기간',D['period']),('매수 밴드',D['band'],'손절가(Stop)',D['stop']),('1차 목표',D['t1'],'2차 목표',D['t2'])]
    for ri,(l1,v1,l2,v2) in enumerate(plan):
        shade(tp.rows[ri].cells[0],LGOLD); shade(tp.rows[ri].cells[2],LGOLD)
        cell_text(tp.rows[ri].cells[0],l1,size=8.5,bold=True,color=GOLDD); cell_text(tp.rows[ri].cells[1],v1,size=9,bold=True,align='center')
        cell_text(tp.rows[ri].cells[2],l2,size=8.5,bold=True,color=GOLDD); cell_text(tp.rows[ri].cells[3],v2,size=9,bold=True,align='center')
    ct=main.add_paragraph(); ct.paragraph_format.space_before=Pt(8); ct.paragraph_format.space_after=Pt(2)
    setfont(ct.add_run('▌ 핵심결론'),size=10,bold=True,color=GOLDD)
    cpp=main.add_paragraph(); cpp.paragraph_format.space_after=Pt(4); cpp.paragraph_format.line_spacing=1.1
    setfont(cpp.add_run(D['call']),size=9.5)
    ip=main.add_paragraph(); ip.paragraph_format.space_before=Pt(6); ip.paragraph_format.space_after=Pt(2)
    setfont(ip.add_run('▌ 투자포인트'),size=10,bold=True,color=GOLDD)
    for n,t in zip(['①','②','③'],D['points1']):
        pp=main.add_paragraph(); pp.paragraph_format.space_after=Pt(2); pp.paragraph_format.line_spacing=1.1
        setfont(pp.add_run(f'{n} '),size=9,bold=True,color=GOLD); setfont(pp.add_run(t),size=9)
    heading_bar(doc,'Forecasts & Valuations  (단위: 억원, 원, %, 배 / E는 컨센서스)',space_before=2,size=11)
    ncol=len(D['fin_years'])+1
    ft=doc.add_table(rows=len(D['fin'])+1,cols=ncol); ft.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(ft,LINE,4)
    w0=4.0; wr=(17.8-w0)/len(D['fin_years']); colw(ft,[w0]+[wr]*len(D['fin_years']))
    for i,y in enumerate(['']+list(D['fin_years'])):
        shade(ft.rows[0].cells[i],GOLDD); cell_text(ft.rows[0].cells[i],y,size=8,bold=True,color='FFFFFF',align='left' if i==0 else 'center')
    for ri,row in enumerate(D['fin'],start=1):
        for ci,val in enumerate(row):
            cell_text(ft.rows[ri].cells[ci],val,size=8,bold=(ci==0),align='left' if ci==0 else 'center')
        if ri%2==0:
            for ci in range(ncol): shade(ft.rows[ri].cells[ci],LIGHT)
    # PAGE 2
    num_heading(doc,'1','기업 개요','Company Snapshot',space_before=0,pb=True)
    cs=doc.add_table(rows=3,cols=4); cs.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(cs,LINE,4); colw(cs,[3.2,5.7,3.2,5.7])
    for ri,(l1,v1,l2,v2) in enumerate(D['snapshot']):
        shade(cs.rows[ri].cells[0],LIGHT); shade(cs.rows[ri].cells[2],LIGHT)
        cell_text(cs.rows[ri].cells[0],l1,size=9,bold=True,color=DGRAY); cell_text(cs.rows[ri].cells[1],v1,size=9)
        cell_text(cs.rows[ri].cells[2],l2,size=9,bold=True,color=DGRAY); cell_text(cs.rows[ri].cells[3],v2,size=9)
    body(doc,D['biz'],size=9,color=DGRAY,space=4,lead='사업 한 줄 요약:  ',lead_color=GOLDD)
    num_heading(doc,'2','매수 이유','Investment Thesis')
    for lead,txt in D['thesis']:
        body(doc,txt,size=9.5,space=3,lead=lead,lead_color=GOLDD)
    num_heading(doc,'3','밸류에이션','Valuation')
    body(doc,D['val_text'],size=9.5,space=4)
    if D.get('peers'):
        subhead(doc,'■ 동종업계 비교 (Peers, LTM 기준)')
        ptb=doc.add_table(rows=len(D['peers'])+1,cols=4); ptb.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(ptb,LINE,4); colw(ptb,[5.0,4.26,4.26,4.26])
        for i,t in enumerate(['종목','PER(배)','PBR(배)','ROE(%)']):
            shade(ptb.rows[0].cells[i],GOLDD); cell_text(ptb.rows[0].cells[i],t,size=9,bold=True,color='FFFFFF',align='center')
        for ri,row in enumerate(D['peers'],start=1):
            emph = row[0].startswith('(') or ('현재' in row[0]) or ('대상' in row[0])
            for ci,v in enumerate(row):
                cell_text(ptb.rows[ri].cells[ci],v,size=9,bold=emph,align='left' if ci==0 else 'center')
            if row[0].startswith('('):
                for ci in range(4): shade(ptb.rows[ri].cells[ci],LIGHT)
        if D.get('peer_note'): body(doc,D['peer_note'],size=8,color=GRAY,italic=True,space=2)
        else:
            sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(2)
    subhead(doc,'■ 목표가 산정 (Target Price)')
    tpv=doc.add_table(rows=2,cols=5); tpv.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(tpv,LINE,4); colw(tpv,[3.56]*5)
    for i,t in enumerate(['적용 이익','목표 멀티플','목표주가','현재주가','상승여력']):
        shade(tpv.rows[0].cells[i],GOLDD); cell_text(tpv.rows[0].cells[i],t,size=8.5,bold=True,color='FFFFFF',align='center')
    for i,(v,em) in enumerate(D['tpcalc']):
        cell_text(tpv.rows[1].cells[i],v,size=9.5 if em else 9,bold=True,color=GOLDD if em else '000000',align='center')
    body(doc,D['val_note'],size=8.5,color=GRAY,italic=True,space=2)
    # PAGE 3
    num_heading(doc,'4','매매 전략 상세','Trading Strategy',space_before=0,pb=True)
    ts=doc.add_table(rows=2,cols=4); ts.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(ts,LINE,4); colw(ts,[4.45]*4)
    for i,t in enumerate(['매수가','손절가','목표가','전략 유형']):
        shade(ts.rows[0].cells[i],GOLDD); cell_text(ts.rows[0].cells[i],t,size=9.5,bold=True,color='FFFFFF',align='center')
    for i,v in enumerate(D['ts_row']):
        cell_text(ts.rows[1].cells[i],v,size=9.5,bold=True,color=GOLDD if i==3 else '000000',align='center')
    body(doc,D['ts_text'],size=9.5,space=3)
    num_heading(doc,'5','리스크 요인 및 재검토 조건','Risks & Review Triggers')
    for t in D['risks']:
        body(doc,t,size=9.5,bullet=True,space=1)
    body(doc,D['review'],size=9.5,space=3,lead='재검토(매도) 조건:  ',lead_color=GOLDD)
    num_heading(doc,'6','결론','Conclusion')
    body(doc,D['concl'],size=9.5,space=4)
    heading_bar(doc,'부록 : 요약 재무제표  (단위: 억원)',size=10)
    ny=len(D['appx_years'])
    ap=doc.add_table(rows=len(D['appx'])+1,cols=ny+1); ap.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(ap,LINE,4)
    colw(ap,[6.0]+[(17.8-6.0)/ny]*ny)
    for i,t in enumerate(['구분']+list(D['appx_years'])):
        shade(ap.rows[0].cells[i],GOLDD); cell_text(ap.rows[0].cells[i],t,size=9,bold=True,color='FFFFFF',align='center')
    for ri,row in enumerate(D['appx'],start=1):
        for ci,val in enumerate(row):
            cell_text(ap.rows[ri].cells[ci],val,size=8.5,bold=(ci==0),align='left' if ci==0 else 'center')
    out=os.path.join(outdir, D['fname']); doc.save(out); print("SAVED:", out); return out

def _norm(D):
    # JSON 로 들어온 tpcalc([[v,em],...]) 를 (v,em) 튜플로 정규화
    if 'tpcalc' in D:
        D['tpcalc']=[(x[0], bool(x[1])) for x in D['tpcalc']]
    if 'thesis' in D:
        D['thesis']=[(x[0], x[1]) for x in D['thesis']]
    if 'stock' in D:
        D['stock']=[(x[0], x[1]) for x in D['stock']]
    return D

if __name__=='__main__':
    if len(sys.argv)<2:
        print("usage: python gen_report.py stocks.json [outdir]"); sys.exit(1)
    with open(sys.argv[1], encoding='utf-8') as f:
        data=json.load(f)
    outdir=sys.argv[2] if len(sys.argv)>2 else OUTDIR
    if isinstance(data, dict): data=[data]
    for D in data:
        build(_norm(D), outdir)
