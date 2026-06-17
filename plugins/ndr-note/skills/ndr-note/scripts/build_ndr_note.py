# -*- coding: utf-8 -*-
"""
기업탐방노트(Nexus Asset Management 양식) 생성기. (OS 독립 / 포터블)
사용법:  python build_ndr_note.py <config.json>
config.json 스키마는 같은 폴더의 config_amotech.json(예시) 참고.
출력: config의 "output" 경로에 .docx 저장.
  - "output"이 절대경로면 그대로 사용.
  - 파일명만/상대경로면 사용자 바탕화면(~/Desktop)에 저장.
  - "output" 생략 시 "<종목명> 탐방 보고서 (Nexus양식).docx" 로 바탕화면에 저장.
PDF 변환은 같은 폴더의 convert_to_pdf.py 사용(LibreOffice 또는 Windows+Word).
"""
import sys, json, io, os

def resolve_output(cfg):
    # 파일명은 항상 "<종목명>_탐방노트.docx" 로 통일한다.
    fname = cfg["cover"]["name"] + "_탐방노트.docx"
    out = cfg.get("output")
    out = os.path.expanduser(out) if out else None
    # output이 절대경로(폴더/파일)면 그 폴더에, 아니면 바탕화면에 통일 파일명으로 저장.
    if out and os.path.isabs(out):
        d = out if os.path.isdir(out) else os.path.dirname(out)
        return os.path.join(d, fname)
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    if not os.path.isdir(desktop):
        desktop = os.path.expanduser("~")
    return os.path.join(desktop, fname)
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- assets ----------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(SCRIPT_DIR, "..", "assets", "nexus_logo.png")

# ---------- palette ----------
GOLD=(0x7A,0x66,0x3C); GOLD_A=(0x9D,0x84,0x4E); BODY=(0x33,0x33,0x33)
GREY=(0x80,0x80,0x80); WHITE=(0xFF,0xFF,0xFF); BLACK=(0x00,0x00,0x00)
GOLD_HEX="7A663C"; BEIGE="F5F2EA"

def setf(run,name="맑은 고딕",size=10,bold=False,color=None,underline=False):
    run.font.size=Pt(size); run.font.bold=bold; run.font.underline=underline
    if color: run.font.color.rgb=RGBColor(*color)
    rPr=run._element.get_or_add_rPr(); rF=rPr.get_or_add_rFonts()
    rF.set(qn('w:eastAsia'),name); rF.set(qn('w:ascii'),name); rF.set(qn('w:hAnsi'),name)

def shade_el(el_pr,fill):
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill); el_pr.append(shd)
def cell_bg(cell,fill): shade_el(cell._tc.get_or_add_tcPr(),fill)
def cell_valign(cell,v="center"):
    tcPr=cell._tc.get_or_add_tcPr(); e=OxmlElement('w:vAlign'); e.set(qn('w:val'),v); tcPr.append(e)

def set_cell(cell,text,size=8.5,bold=False,color=BODY,align=None,bg=None,valign="center"):
    cell.text=""; p=cell.paragraphs[0]
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)
    r=p.add_run(text); setf(r,size=size,bold=bold,color=color)
    if bg: cell_bg(cell,bg)
    cell_valign(cell,valign); return p

def _borders(table,val,color):
    tblPr=table._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),val); e.set(qn('w:sz'),'4' if val=='single' else '0'); e.set(qn('w:space'),'0')
        if val=='single': e.set(qn('w:color'),color)
        b.append(e)
    tblPr.append(b)
def no_borders(t): _borders(t,'none','auto')
def thin_borders(t,color="DDD6C4"): _borders(t,'single',color)

def set_widths(table,widths):
    table.autofit=False
    for row in table.rows:
        for i,c in enumerate(row.cells): c.width=Cm(widths[i])
    for i,col in enumerate(table.columns): col.width=Cm(widths[i])

# global doc + helpers bound at runtime
DOC=None
def para(space_after=4,space_before=0,align=None,indent=None,line=None,container=None):
    p=(container.add_paragraph() if container is not None else DOC.add_paragraph())
    pf=p.paragraph_format; pf.space_after=Pt(space_after); pf.space_before=Pt(space_before)
    if align is not None: p.alignment=align
    if indent is not None: pf.left_indent=Cm(indent)
    if line: pf.line_spacing=line
    return p

def rich_md(p, text, size=9.5, color=BODY, emph_underline=True):
    """**...** 구간을 볼드(+밑줄)로 렌더링."""
    parts=text.split("**")
    for i,seg in enumerate(parts):
        if seg=="": continue
        bold=(i%2==1)
        r=p.add_run(seg); setf(r,size=size,bold=bold,color=color,underline=(bold and emph_underline))
    return p

def hrule(p,color=GOLD_HEX,sz="18"):
    pPr=p._element.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr')
    bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),sz); bottom.set(qn('w:space'),'2'); bottom.set(qn('w:color'),color)
    pbdr.append(bottom); pPr.append(pbdr)

def keep_next(p):
    pPr=p._p.get_or_add_pPr()
    if pPr.find(qn('w:keepNext')) is None: pPr.append(OxmlElement('w:keepNext'))
def cant_split(row):
    trPr=row._tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None: trPr.append(OxmlElement('w:cantSplit'))
def keep_table_together(tbl, heading_p=None):
    if heading_p is not None: keep_next(heading_p)
    rows=tbl.rows
    for ri,row in enumerate(rows):
        cant_split(row)
        if ri<len(rows)-1:
            for c in row.cells:
                for p in c.paragraphs: keep_next(p)

def mini_header(tbl,ncols,text):
    m=tbl.rows[0].cells[0].merge(tbl.rows[0].cells[ncols-1])
    set_cell(m,text,size=8.5,bold=True,color=WHITE,align=WD_ALIGN_PARAGRAPH.CENTER,bg=GOLD_HEX)

def build(cfg):
    global DOC
    DOC=Document()
    s=DOC.sections[0]
    s.page_width=Emu(7772400); s.page_height=Emu(10058400)
    s.left_margin=Emu(539750); s.right_margin=Emu(539750); s.top_margin=Emu(431800); s.bottom_margin=Emu(431800)
    n=DOC.styles['Normal']; n.font.name="맑은 고딕"; n.font.size=Pt(10)
    n.element.rPr.rFonts.set(qn('w:eastAsia'),"맑은 고딕")

    meta=cfg["meta"]; cov=cfg["cover"]; rb=cfg.get("reported_by",{})

    # ---- header band ----
    hb=DOC.add_table(rows=1,cols=2); no_borders(hb); set_widths(hb,[9.5,9.1])
    lc=hb.rows[0].cells[0]
    if os.path.isfile(LOGO_PATH):
        # 로고 이미지 삽입 (없으면 회사명 텍스트로 폴백)
        lc.text=""; lp=lc.paragraphs[0]; lp.alignment=WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_after=Pt(0); lp.paragraph_format.space_before=Pt(0)
        lp.add_run().add_picture(LOGO_PATH, height=Cm(0.85)); cell_valign(lc,"center")
    else:
        set_cell(lc,meta["firm"],size=11,bold=True,color=GOLD,align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(hb.rows[0].cells[1],"%s   |   %s"%(meta["note_type"],meta["date"]),size=9,color=GREY,align=WD_ALIGN_PARAGRAPH.RIGHT)
    hrule(para(space_after=4))

    # ---- body 2-col ----
    body=DOC.add_table(rows=1,cols=2); no_borders(body); set_widths(body,[5.7,12.9])
    L=body.rows[0].cells[0]; R=body.rows[0].cells[1]; L.text=""; R.text=""

    # Stock Data
    sd=cov["stock_data"]
    t2=L.add_table(rows=len(sd)+1,cols=2); thin_borders(t2); set_widths(t2,[2.5,3.0])
    mini_header(t2,2,"Stock Data")
    for i,(k,v) in enumerate(sd):
        set_cell(t2.rows[i+1].cells[0],k,size=8,bold=True,color=BODY,bg=BEIGE)
        set_cell(t2.rows[i+1].cells[1],v,size=8,color=BODY,align=WD_ALIGN_PARAGRAPH.RIGHT)
    L.add_paragraph().paragraph_format.space_after=Pt(3)

    # 상대수익률
    t3=L.add_table(rows=3,cols=5); thin_borders(t3); set_widths(t3,[1.5,1.0,1.0,1.0,1.0])
    mini_header(t3,5,"상대수익률 (KOSPI 대비, %p)")
    for j,h in enumerate(["","1M","3M","6M","12M"]):
        set_cell(t3.rows[1].cells[j],h,size=7.5,bold=True,color=BODY,align=WD_ALIGN_PARAGRAPH.CENTER,bg=BEIGE)
    rel=cov.get("rel_return",["–","–","–","–"])
    for j,v in enumerate(["초과"]+rel):
        set_cell(t3.rows[2].cells[j],v,size=7.5,bold=(j==0),color=BODY,align=WD_ALIGN_PARAGRAPH.CENTER)
    L.add_paragraph().paragraph_format.space_after=Pt(3)

    # 주가추이
    t4=L.add_table(rows=2,cols=1); thin_borders(t4); set_widths(t4,[5.5])
    set_cell(t4.rows[0].cells[0],"주가 추이 (12M)",size=8.5,bold=True,color=WHITE,align=WD_ALIGN_PARAGRAPH.CENTER,bg=GOLD_HEX)
    c=t4.rows[1].cells[0]; set_cell(c,"[ 주가/상대강도 차트 삽입 ]",size=8,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER)
    c.paragraphs[0].paragraph_format.space_before=Pt(8); c.paragraphs[0].paragraph_format.space_after=Pt(8)
    L.add_paragraph().paragraph_format.space_after=Pt(4)

    # reported by
    p=para(space_after=0,container=L); r=p.add_run("Reported by"); setf(r,size=7.5,color=GREY)
    p=para(space_after=0,container=L)
    r=p.add_run((rb.get("name","")+"  ")); setf(r,size=8.5,bold=True,color=GOLD)
    r=p.add_run(rb.get("role","")); setf(r,size=8.5,bold=True,color=GOLD)
    if rb.get("email"):
        p=para(space_after=0,container=L); r=p.add_run(rb["email"]); setf(r,size=7.5,color=GREY)

    # ===== right =====
    p=para(space_after=0,container=R)
    r=p.add_run(cov["name"]+" "); setf(r,size=17,bold=True,color=GOLD)
    r=p.add_run("("+cov["code"]+")"); setf(r,size=12,bold=True,color=BODY)
    p=para(space_after=4,container=R); r=p.add_run(cov["sector"]); setf(r,size=8.5,color=GREY)
    p=para(space_after=6,container=R); r=p.add_run(cov["headline"]); setf(r,size=13,bold=True,color=BLACK)

    # ▌핵심 요약
    p=para(space_after=1,space_before=2,container=R); r=p.add_run("▌ 핵심 요약"); setf(r,size=10,bold=True,color=GOLD)
    rich_md(para(space_after=3,container=R,line=1.15),cov["summary"],size=8.5,color=BODY)

    # ▌탐방 포인트
    p=para(space_after=1,space_before=1,container=R); r=p.add_run("▌ 탐방 포인트"); setf(r,size=10,bold=True,color=GOLD)
    marks="①②③④⑤"
    for i,(h,t) in enumerate(cov["points"]):
        pp=para(space_after=1,container=R,indent=0.1,line=1.1)
        r=pp.add_run(marks[i]+" "); setf(r,size=8.5,bold=True,color=GOLD_A)
        r=pp.add_run(h+" — "); setf(r,size=8.5,bold=True,color=GOLD)
        r=pp.add_run(t); setf(r,size=8.5,color=BODY)

    # Forecasts & Valuations
    fc=cfg["forecasts"]
    pfv=para(space_after=2,space_before=4); r=pfv.add_run("Forecasts & Valuations  "); setf(r,size=11,bold=True,color=GOLD)
    r=pfv.add_run(fc.get("unit","(단위: 억원, 원, %, 배 / E는 자체 추정)")); setf(r,size=8,color=GREY)
    yrs=fc["years"]; rows=fc["rows"]
    fv=DOC.add_table(rows=len(rows)+1,cols=len(yrs)+1); thin_borders(fv)
    w=[4.0]+[ (15.0/len(yrs)) for _ in yrs]; set_widths(fv,w)
    set_cell(fv.rows[0].cells[0],"",bg=GOLD_HEX)
    for j,y in enumerate(yrs): set_cell(fv.rows[0].cells[j+1],y,size=8.5,bold=True,color=WHITE,align=WD_ALIGN_PARAGRAPH.CENTER,bg=GOLD_HEX)
    for i,row in enumerate(rows):
        bg=BEIGE if i%2==1 else None
        set_cell(fv.rows[i+1].cells[0],row[0],size=8.5,bold=True,color=BODY,bg=bg)
        for j,v in enumerate(row[1:]):
            set_cell(fv.rows[i+1].cells[j+1],v,size=8.5,color=BODY,align=WD_ALIGN_PARAGRAPH.CENTER,bg=bg)
    keep_table_together(fv,pfv)
    if fc.get("footnote"):
        p=para(space_before=1); r=p.add_run("※ "+fc["footnote"]); setf(r,size=7.5,color=GREY)

    # ===== PAGE 2: deep dive =====
    dd=cfg["deepdive"]
    DOC.add_page_break()
    p=para(space_before=4,space_after=8,align=WD_ALIGN_PARAGRAPH.CENTER); r=p.add_run(dd["title"]); setf(r,size=15,bold=True,color=GOLD)
    for sec in dd["sections"]:
        p=para(space_before=8,space_after=3); r=p.add_run("▌ "+sec["head"]); setf(r,size=10.5,bold=True,color=GOLD)
        rich_md(para(space_after=6,line=1.25),sec["body"])
    rm=dd.get("roadmap")
    if rm:
        p=para(space_before=8,space_after=3); r=p.add_run("▌ "+rm.get("title","제품/사업 로드맵")); setf(r,size=10.5,bold=True,color=GOLD)
        hdr=rm["header"]; rws=rm["rows"]
        rt=DOC.add_table(rows=len(rws)+1,cols=len(hdr)); thin_borders(rt)
        set_widths(rt,rm.get("widths",[18.6/len(hdr)]*len(hdr)))
        for j,h in enumerate(hdr): set_cell(rt.rows[0].cells[j],h,size=8.5,bold=True,color=WHITE,align=WD_ALIGN_PARAGRAPH.CENTER,bg=GOLD_HEX)
        for i,row in enumerate(rws):
            for j,v in enumerate(row):
                set_cell(rt.rows[i+1].cells[j],v,size=8,bold=(j==0),color=BODY,bg=(BEIGE if i%2==1 else None),
                         align=(WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER))
        keep_table_together(rt)
    if dd.get("source"):
        p=para(space_before=2); r=p.add_run(dd["source"]); setf(r,size=7.5,color=GREY)

    # ===== PAGE 3: Q&A =====
    DOC.add_page_break()
    qp=para(space_after=6); shade_el(qp._element.get_or_add_pPr(),GOLD_HEX)
    r=qp.add_run("  Q&A  "); setf(r,size=13,bold=True,color=WHITE)
    for q,a in cfg["qa"]:
        pq=para(space_before=6,space_after=2)
        r=pq.add_run("Q.  "); setf(r,size=9.5,bold=True,color=GOLD_A)
        r=pq.add_run(q); setf(r,size=9.5,bold=True,color=GOLD,underline=True)
        pa=para(space_after=3,indent=0.55,line=1.2)
        r=pa.add_run("A.  "); setf(r,size=9.5,bold=True,color=BODY)
        r=pa.add_run(a); setf(r,size=9.5,color=BODY)

    out=resolve_output(cfg)
    os.makedirs(os.path.dirname(out), exist_ok=True)
    DOC.save(out); return out

if __name__=="__main__":
    if len(sys.argv)<2:
        print("사용법: python build_ndr_note.py <config.json>"); sys.exit(1)
    cfg_path=sys.argv[1]
    with io.open(cfg_path,"r",encoding="utf-8") as f: cfg=json.load(f)
    out=build(cfg)
    print("SAVED:",out)
